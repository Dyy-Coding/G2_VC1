import sys
import json
from docx import Document

def main():
    if len(sys.argv) < 3:
        print("Usage: python export_word.py input.json output.docx")
        sys.exit(1)

    json_path = sys.argv[1]
    output_path = sys.argv[2]

    with open(json_path, 'r') as f:
        materials = json.load(f)

    doc = Document()
    doc.add_heading('Materials List', level=1)

    for material in materials:
        for key, value in material.items():
            doc.add_paragraph(f"{key}: {value}")
        doc.add_paragraph('')  # Empty line

    doc.save(output_path)

if __name__ == "__main__":
    main()
